from docx import Document
import csv
import json
import re
import pandas as pd
from pathlib import Path

def extract_table(index, outfile):
    doc = Document("data/input/A_80_400.DOCX")
    with open(outfile, "w", newline="", encoding="utf-8") as f:
        writer = csv.writer(f)
        for row in doc.tables[index].rows:
            writer.writerow([cell.text.strip() for cell in row.cells])


def verify_hierarchy(df, grand_total=None):
    numeric_cols = [c for c in df.select_dtypes('number').columns if c != 'Section']
    # exclude percentage columns from grand total check (can't sum percentages)
    summable_cols = [c for c in numeric_cols if 'percentage' not in c.lower()]
    errors = []
    # check grand total = sum of part totals
    if grand_total is not None:
        part_sum = df[df['row_type'] == 'part_total'][summable_cols].sum()
        diff = (grand_total[summable_cols] - part_sum).abs()
        if (diff > 0.1).any():
            errors.append(f"Grand total != sum(part_totals), diff={diff[diff > 0.1].to_dict()}")
    # check part totals = sum of section totals
    for part in df['Part'].unique():
        part_df = df[df['Part'] == part]
        part_total = part_df[part_df['row_type'] == 'part_total'][numeric_cols].iloc[0]
        section_sum = part_df[part_df['row_type'] == 'section_total'][numeric_cols].sum()
        if len(part_df[part_df['row_type'] == 'section_total']) > 0:
            diff = (part_total - section_sum).abs()
            if (diff > 0.1).any():
                errors.append(f"{part}: part_total != sum(section_totals), diff={diff[diff > 0.1].to_dict()}")
    # check section totals = sum of entity totals
    for (part, section), sec_df in df.groupby(['Part', 'Section']):
        if pd.isna(section):
            continue
        entities = sec_df[sec_df['row_type'] == 'entity_total']
        section_totals = sec_df[sec_df['row_type'] == 'section_total']
        if len(entities) == 0 or len(section_totals) == 0:
            continue
        section_total = section_totals[numeric_cols].iloc[0]
        entity_sum = entities[numeric_cols].sum()
        diff = (section_total - entity_sum).abs()
        if (diff > 0.1).any():
            errors.append(f"{part} Section {section}: section_total != sum(entity_totals), diff={diff[diff > 0.1].to_dict()}")
    return errors

def clean_table6():
    # the table contains 3 row types
    # - part totals (stretched over 2 rows)
    # - section totals (=breakdowns of parts)
    # - entity totals (=breakdowns of sections, but not available for every section)
    # additionally, for section 3 there is a weird "– Other" category which seems to be the total of the subsequent rows and is thus redundant, so we drop that row
    df = pd.read_csv("data/intermediate/table6_raw.csv", header=[0, 1], skiprows=[2])
    df.columns = [a if a == b or pd.isna(b) or b == '' else f'{a} – {b}' for a, b in df.columns]
    col = 'Budget part/section/entity'
    is_part_header = df[col].str.startswith('Part ')
    is_part_total = is_part_header.shift(1, fill_value=False)
    is_section_total = df[col].str.match(r'^\d+[A-Z]?\.\t')
    is_entity_total = ~is_part_header & ~is_part_total & ~is_section_total
    df['row_type'] = 'entity_total'
    df.loc[is_part_total, 'row_type'] = 'part_total'
    df.loc[is_section_total, 'row_type'] = 'section_total'
    df['Part'] = df[col].where(is_part_header).ffill()
    df['Part name'] = df[col].where(is_part_total).ffill()
    df[['Section', 'Section name']] = df[col].str.extract(r'^(\d+[A-Z]?)\.\t(.+)')
    df[['Section', 'Section name']] = df.groupby('Part')[['Section', 'Section name']].ffill()
    df['Entity name'] = df[col].where(is_entity_total).str.replace(r'^–\t', '', regex=True)
    # special case: "Resident coordinator systema" has footnote 'a' attached
    df['_entity_fn'] = df['Entity name'].where(df['Entity name'] == 'Resident coordinator systema').str.extract(r'([a-z])$', expand=False)
    df['Entity name'] = df['Entity name'].replace('Resident coordinator systema', 'Resident coordinator system')
    order = ['row_type', 'Part', 'Part name', 'Section', 'Section name', 'Entity name']
    numeric_cols = [c for c in df.columns if c not in order + [col, '_entity_fn']]
    # extract footnote markers before cleaning (from numeric cols + entity name)
    df['footnotes'] = df[numeric_cols].apply(lambda row: ''.join(sorted(set(''.join(row.dropna().astype(str).str.extract(r'([a-z]+)$', expand=False).dropna())))), axis=1)
    df['footnotes'] = (df['_entity_fn'].fillna('') + df['footnotes']).apply(lambda x: ''.join(sorted(set(x))) or pd.NA)
    # clean numeric columns for ALL rows first (values are in thousands in source)
    for c in numeric_cols:
        df[c] = df[c].replace('–', '0').str.replace(r'[\s,]', '', regex=True).str.replace(r'\((.+)\)[a-z]*', r'-\1', regex=True).str.replace(r'[a-z]+$', '', regex=True)
        df[c] = pd.to_numeric(df[c], errors='coerce')
        if 'percentage' not in c.lower():
            df[c] = df[c] * 1000  # convert from thousands to actual values
    # extract grand total after cleaning, before filtering
    grand_total = df[df[col] == 'Total'][numeric_cols].iloc[0]
    df = df[~is_part_header & ~df[col].isin(['–\tOther', 'Total']) & df['2025 approved'].notna()]
    df = df[order + numeric_cols + ['footnotes']]

    # save grand total as a separate row
    grand_total_row = pd.DataFrame([{
        'row_type': 'grand_total',
        'Part': None,
        'Part name': 'Total',
        'Section': None,
        'Section name': None,
        'Entity name': None,
        **grand_total.to_dict(),
        'footnotes': None,
    }])
    
    # extraction is finished now, but there are errors
    errors = verify_hierarchy(df, grand_total)
    assert len(errors) == 2  # known data errors in source: Transitional capacities not aggregated
    # manual fix: roll up Transitional capacities through hierarchy
    tc = 'Transitional capacities'
    # fix section totals from entity sums (only if entities have TC values)
    for (part, section), grp in df.groupby(['Part', 'Section']):
        if pd.isna(section):
            continue
        entity_sum = grp[grp['row_type'] == 'entity_total'][tc].sum()
        if entity_sum > 0:
            df.loc[(df['Part'] == part) & (df['Section'] == section) & (df['row_type'] == 'section_total'), tc] = entity_sum
    # fix part totals from section sums
    for part in df['Part'].unique():
        section_sum = df[(df['Part'] == part) & (df['row_type'] == 'section_total')][tc].sum()
        df.loc[(df['Part'] == part) & (df['row_type'] == 'part_total'), tc] = section_sum
    errors = verify_hierarchy(df, grand_total)
    assert not errors
    # prepend grand total row to the dataframe
    df = pd.concat([grand_total_row, df], ignore_index=True)
    df.to_csv("data/intermediate/table6.csv", index=False)
    return df


def parse_paragraph(text):
    """Parse paragraph prefix and determine hierarchy level."""
    patterns = [
        (r'^(\d+)\.\t(.+)$', 0),           # "92.\t..." → level 0
        (r'^\(([ivx]+)\)\t(.+)$', 2),      # "(i)\t..." → level 2 (roman numerals, check first)
        (r'^\(([a-z])\)\t(.+)$', 1),       # "(a)\t..." → level 1 (letters)
    ]
    for pat, level in patterns:
        m = re.match(pat, text, re.DOTALL)
        if m:
            return {'prefix': m.group(1), 'level': level, 'text': m.group(2)}
    return {'prefix': None, 'level': None, 'text': text}

def extract_resource_table(tbl):
    """Extract resource changes table data."""
    if len(tbl.rows) < 4:
        return None
    headers = [c.text.strip() for c in tbl.rows[1].cells]
    rows = []
    in_transitional = False
    for row in tbl.rows[3:]:
        cells = [c.text.strip() for c in row.cells]
        label = cells[0].lower() if cells[0] else ''
        # Skip empty rows, variance rows, subtotals, and totals
        if not cells[0] or 'variance' in label or 'subtotal' in label or label == 'total':
            continue
        # Check for transitional section header (has empty numeric values)
        if 'transitional' in label and all(c in ['', '–', '0'] for c in cells[1:]):
            in_transitional = True
            continue
        # Clean numeric values
        clean_label = f"Transitional: {cells[0]}" if in_transitional else cells[0]
        clean = [clean_label]
        for v in cells[1:]:
            v = v.replace('\xa0', '').replace(' ', '').replace(',', '').replace('–', '0')
            v = re.sub(r'\((.+)\)', r'-\1', v)  # (123) → -123
            clean.append(v)
        rows.append(clean)
    return {'headers': headers, 'rows': rows} if rows else None

def extract_entity_details():
    """Extract narrative paragraphs and resource tables for each entity in Chapter IV.B."""
    from docx.oxml.ns import qn
    doc = Document("data/input/A_80_400.DOCX")
    body = doc.element.body
    
    # Build sequence of body elements
    elements = []
    para_idx, table_idx = 0, 0
    for elem in body:
        if elem.tag == qn('w:p'):
            elements.append(('para', para_idx, doc.paragraphs[para_idx].text.strip(), doc.paragraphs[para_idx].style.name))
            para_idx += 1
        elif elem.tag == qn('w:tbl'):
            elements.append(('table', table_idx, '', ''))
            table_idx += 1
    
    # Find entity sections and their content
    entity_starts = []
    for i, (typ, idx, text, style) in enumerate(elements):
        if typ == 'para':
            m = re.match(r'^(\d+)\.\s+Section\s+(\d+[A-Z]?),\s+(.+)$', text)
            if m:
                entity_starts.append((i, int(m.group(1)), m.group(2), m.group(3).strip()))
    
    sections = []
    for j, (start_i, num, sec, name) in enumerate(entity_starts):
        end_i = entity_starts[j+1][0] if j+1 < len(entity_starts) else len(elements)
        entry = {'num': num, 'section': sec, 'entity': name, 'narratives': [], 'resource_table': None}
        
        for i in range(start_i, end_i):
            typ, idx, text, style = elements[i]
            if typ == 'para':
                # Extract narratives
                if style == '__Single Txt' and len(text) > 50:
                    entry['narratives'].append(parse_paragraph(text))
                # Find "Regular budget: proposed resource changes" table
                if 'Regular budget: proposed resource changes by object of expenditure' in text:
                    # Next table element is the resource table
                    for k in range(i+1, min(i+5, end_i)):
                        if elements[k][0] == 'table':
                            tbl = doc.tables[elements[k][1]]
                            first_cell = tbl.rows[0].cells[0].text.strip() if tbl.rows else ''
                            if 'Object of expenditure' in first_cell or not first_cell:
                                entry['resource_table'] = extract_resource_table(tbl)
                            break
        sections.append(entry)
    
    with open("data/intermediate/entity_narratives.json", "w", encoding="utf-8") as f:
        json.dump(sections, f, indent=2, ensure_ascii=False)
    return sections

# budget overview based on table 6
extract_table(9, "data/intermediate/table6_raw.csv")
budget = clean_table6()
name_map = pd.read_csv("data/input/entity_name_mapping.csv").set_index('table6_name')['doc_name'].to_dict()
# Map Entity name first, then fall back to Section name for section-only rows
budget['chapter_title'] = budget.apply(
    lambda r: name_map.get(r['Entity name'], r['Entity name']) if pd.notna(r['Entity name']) 
              else name_map.get(r['Section name'], r['Section name']), axis=1
)

# Load UN entities for abbreviations and entity names
with open("data/input/2025-12-05_un-entities.json", encoding="utf-8") as f:
    un_entities = json.load(f)
# Create mapping: entity_long -> (abbreviation, entity_long)
entity_info = {e['entity_long']: (e['entity'], e['entity_long']) for e in un_entities if e.get('entity') and e.get('entity_long')}
# Also add direct matches by abbreviation (entity field)
entity_by_abbr = {e['entity']: (e['entity'], e['entity_long']) for e in un_entities if e.get('entity') and e.get('entity_long')}

def get_entity_info(title):
    """Return (abbreviation, entity_name) for a chapter_title."""
    if not title:
        return (None, None)
    # Direct match by full name
    if title in entity_info:
        return entity_info[title]
    # Direct match by abbreviation (for titles like "UN-Habitat")
    if title in entity_by_abbr:
        return entity_by_abbr[title]
    return (None, None)

budget['abbreviation'] = budget['chapter_title'].apply(lambda t: get_entity_info(t)[0])
budget['entity_name'] = budget['chapter_title'].apply(lambda t: get_entity_info(t)[1])
budget.to_json("public/budget.json", orient="records", indent=2, force_ascii=False)

# budget details based on later sections in the document
paras = extract_entity_details()
json.dump(paras, Path("public/details.json").open("w"), indent=2, ensure_ascii=False)
